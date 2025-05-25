import streamlit as st
from paddleocr import PaddleOCR
import fitz  # PyMuPDF
import docx
import tempfile
import os
from PIL import Image
import numpy as np
import re
import heapq
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Initialize OCR
ocr = PaddleOCR(use_angle_cls=True, lang='en')

# OCR for image
def extract_text_from_image(file_path):
    result = ocr.ocr(file_path, cls=True)
    text = []
    for line in result[0]:
        text.append(line[1][0])
    return "\n".join(text)

# Extract text from PDF
def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    extracted_text = []
    for page in doc:
        extracted_text.append(page.get_text("text"))
        image_list = page.get_images(full=True)
        for img in image_list:
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_file:
                img_file.write(image_bytes)
                img_text = extract_text_from_image(img_file.name)
                extracted_text.append(img_text)
            os.unlink(img_file.name)
    return "\n".join(extracted_text)

# Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

# Streamlit UI
st.set_page_config(page_title="RAG Chatbot with PaddleOCR", layout="wide")
st.title("📘 Document Q&A Chatbot (OCR, Tables, Images, Links)")
uploaded_files = st.file_uploader("📤 Upload your files (PDF, DOCX, JPG, PNG)", type=["pdf", "docx", "png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded_files:
    st.success("✅ Files uploaded successfully!")
    full_corpus = ""
    for uploaded_file in uploaded_files:
        suffix = uploaded_file.name.split(".")[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name
        if suffix == "pdf":
            full_corpus += extract_text_from_pdf(tmp_path)
        elif suffix == "docx":
            full_corpus += extract_text_from_docx(tmp_path)
        elif suffix in ["jpg", "png", "jpeg"]:
            full_corpus += extract_text_from_image(tmp_path)
        os.unlink(tmp_path)

    st.text_area("📄 Extracted Document Content", value=full_corpus, height=300)

    question = st.text_input("💬 Ask your question from the documents:")
    if question:
        sentences = [s.strip() for s in re.split(r'(?<=[.!?]) +', full_corpus) if len(s.strip()) > 20]
        vectorizer = TfidfVectorizer().fit_transform(sentences + [question])
        similarity = cosine_similarity(vectorizer[-1], vectorizer[:-1])
        top_n = heapq.nlargest(3, zip(similarity[0], sentences), key=lambda x: x[0])
        st.subheader("🤖 Answer")
        for score, sent in top_n:
            st.markdown(f"- {sent}")
